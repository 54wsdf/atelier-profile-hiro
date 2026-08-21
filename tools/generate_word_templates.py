from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'word'
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / 'assets' / 'hiro2026-logo.png'

# Front matter follows AtelierTeX; body follows HIRO colors.
AT_INK='202124'; AT_MUTED='666B73'; AT_RULE='C9CDD4'; AT_ACCENT='334E68'
HIRO_INK='111111'; HIRO_TEXT='1D1D1B'; HIRO_GRAY='666662'; HIRO_LIGHT='969691'; HIRO_RULE='C6C6C0'
SERIF_LATIN='Noto Serif'; SERIF_CJK='Noto Serif CJK SC'
SANS_LATIN='Noto Sans'; SANS_CJK='Noto Sans CJK SC'

META = {
    'paper_type': '第三方支援样张',
    'paper_id': 'HIRO2026-DEMO',
    'short_title': 'HIRO2026 非官方支援模板',
    'kicker': '篠泽广 / 第三方支援项目',
    'title': 'HIRO2026 非官方支援排版模板',
    'subtitle': '维护者个人发起，基于 AtelierTeX 制作',
    'english_title': 'A Personal, Unofficial HIRO2026 Support Profile Built on AtelierTeX',
    'author': '54wsdf / 个人支援项目',
    'affiliation': '面向 HIRO2026 的非官方排版与投稿支援',
    'contact': 'https://idol-master.top/sites/hiro2026',
    'author_note': '本项目纯属个人制作，不是 HIRO2026 官方模板，与活动主办方无隶属、委托、合作或代表关系。活动规则以活动页面为准。',
    'deck': '随笔刊页把阅读入口、作者信息与正文开篇放在同一页面，同时继续使用 HIRO2026 的刊头、章节、图表和参考文献体系。',
}
ABSTRACT = ('本样张展示维护者个人发起、面向篠泽广研讨会（HIRO2026）的第三方非官方 LaTeX 支援模板。'
            '模板建立在 AtelierTeX 的多语种长文框架与 ATX-ACGN-REF 跨媒介引用体系之上，为 HIRO2026 主题写作提供刊头、题名区、章节层级、叙事环境与参考文献视觉。'
            '本项目与 HIRO2026 官方及活动主办方没有关系；活动日程、投稿规则与最终要求始终以活动网站公布的信息为准。')
KEYWORDS = '篠泽广；学园偶像大师；HIRO2026；AtelierTeX；ACGN；LaTeX'


def rpr_fonts(rpr, latin, east_asia):
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:cs'), latin)
    rfonts.set(qn('w:eastAsia'), east_asia)


def set_run_font(run, latin=SERIF_LATIN, east_asia=SERIF_CJK, size=None,
                 bold=None, italic=None, color=None):
    run.font.name = latin
    rpr_fonts(run._element.get_or_add_rPr(), latin, east_asia)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def style_font(style, latin, east_asia, size, bold=None, italic=None, color=None):
    style.font.name = latin
    rpr_fonts(style.element.get_or_add_rPr(), latin, east_asia)
    style.font.size = Pt(size)
    if bold is not None: style.font.bold = bold
    if italic is not None: style.font.italic = italic
    if color: style.font.color.rgb = RGBColor.from_string(color)


def add_pstyle(doc, name, latin=SERIF_LATIN, east_asia=SERIF_CJK, size=10,
               bold=None, italic=None, color=HIRO_TEXT, left=None, right=None,
               first=None, before=None, after=None, line=None, align=None):
    try: s = doc.styles[name]
    except KeyError: s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style_font(s, latin, east_asia, size, bold, italic, color)
    pf = s.paragraph_format
    if left is not None: pf.left_indent = Mm(left)
    if right is not None: pf.right_indent = Mm(right)
    if first is not None: pf.first_line_indent = Mm(first)
    if before is not None: pf.space_before = Pt(before)
    if after is not None: pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if align is not None: pf.alignment = align
    return s


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for tag, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + tag))
        if node is None:
            node = OxmlElement('w:' + tag); tcMar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, mm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:type'), 'dxa'); tcW.set(qn('w:w'), str(int(mm / 25.4 * 1440)))


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn('w:' + edge))
        if el is None:
            el = OxmlElement('w:' + edge); borders.append(el)
        el.set(qn('w:val'), 'nil')


def paragraph_bottom_border(p, color=HIRO_RULE, size='3', space='1'):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), space); bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


def keep_with_next(p, enabled=True):
    p.paragraph_format.keep_with_next = enabled


def add_page_field(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); set_run_font(r, size=7.2, color=HIRO_GRAY)
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
    r._r.extend([fld1, instr, fld2])


def setup_styles(doc):
    normal = doc.styles['Normal']
    style_font(normal, SERIF_LATIN, SERIF_CJK, 10.35, color=HIRO_TEXT)
    normal.paragraph_format.first_line_indent = Mm(7.3)
    normal.paragraph_format.line_spacing = Pt(16.1)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.space_after = Pt(0)

    h1 = doc.styles['Heading 1']
    style_font(h1, SANS_LATIN, SANS_CJK, 14.2, bold=True, color=HIRO_INK)
    h1.paragraph_format.space_before = Pt(24); h1.paragraph_format.space_after = Pt(9)
    h1.paragraph_format.line_spacing = Pt(20); h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h1.paragraph_format.keep_with_next = True
    h2 = doc.styles['Heading 2']
    style_font(h2, SANS_LATIN, SANS_CJK, 11.1, bold=True, color=HIRO_INK)
    h2.paragraph_format.space_before = Pt(16); h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = Pt(16.5); h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h2.paragraph_format.keep_with_next = True

    specs = [
      ('AT Feature Kicker',SANS_LATIN,SANS_CJK,9,True,None,AT_ACCENT,None,None,None,0,11,11,None),
      ('AT Feature Title',SERIF_LATIN,SERIF_CJK,20.7,True,None,AT_INK,None,None,None,0,0,25,None),
      ('AT Feature Subtitle',SERIF_LATIN,SERIF_CJK,14.4,None,None,AT_MUTED,None,None,None,0,0,18,None),
      ('AT Feature English',SERIF_LATIN,SERIF_CJK,12,None,True,AT_MUTED,None,None,None,0,0,15,None),
      ('AT Feature Author',SERIF_LATIN,SERIF_CJK,12,None,None,AT_INK,None,None,None,0,0,15,None),
      ('AT Feature Affiliation',SERIF_LATIN,SERIF_CJK,10,None,None,AT_MUTED,None,None,None,0,0,13,None),
      ('AT Feature Contact',SERIF_LATIN,SERIF_CJK,9,None,None,AT_MUTED,None,None,None,0,0,12,None),
      ('AT Feature Meta',SERIF_LATIN,SERIF_CJK,9,None,None,AT_MUTED,None,None,None,0,0,12,None),
      ('AT Feature Note',SERIF_LATIN,SERIF_CJK,8,None,None,AT_MUTED,None,None,None,0,0,11,None),
      ('AT Pub Meta',SANS_LATIN,SANS_CJK,6.8,None,None,AT_MUTED,None,None,None,0,0,8.5,WD_ALIGN_PARAGRAPH.RIGHT),
      ('AT Kicker',SANS_LATIN,SANS_CJK,7.3,True,None,AT_MUTED,None,None,None,0,0,9,None),
      ('AT Title',SERIF_LATIN,SERIF_CJK,27,True,None,AT_INK,None,None,None,0,0,35,None),
      ('AT Subtitle',SERIF_LATIN,SERIF_CJK,12.6,None,None,AT_INK,None,None,None,0,0,19.2,None),
      ('AT English',SERIF_LATIN,SERIF_CJK,9.7,None,True,AT_MUTED,None,18.5,None,0,0,13.5,None),
      ('AT Deck',SERIF_LATIN,SERIF_CJK,9.7,None,None,AT_MUTED,None,11,None,0,0,15.5,None),
      ('AT Author',SERIF_LATIN,SERIF_CJK,10.1,None,None,AT_INK,None,None,None,0,0,13.4,None),
      ('AT Essay Author',SERIF_LATIN,SERIF_CJK,9.9,None,None,AT_INK,None,None,None,0,0,13.2,None),
      ('AT Affiliation',SERIF_LATIN,SERIF_CJK,8.8,None,None,AT_MUTED,None,None,None,0,0,12.2,None),
      ('AT Essay Affiliation',SERIF_LATIN,SERIF_CJK,8.5,None,None,AT_MUTED,None,None,None,0,0,11.8,None),
      ('AT Contact',SERIF_LATIN,SERIF_CJK,8.5,None,None,AT_MUTED,None,None,None,0,0,11.6,None),
      ('AT Document Meta',SANS_LATIN,SANS_CJK,6.9,None,None,AT_MUTED,None,None,None,0,0,8.6,WD_ALIGN_PARAGRAPH.RIGHT),
      ('AT Note',SERIF_LATIN,SERIF_CJK,7.8,None,None,AT_MUTED,None,None,None,0,0,10.5,None),
      ('HIRO Abstract',SERIF_LATIN,SERIF_CJK,9.45,None,None,HIRO_TEXT,7,7,None,0,0,15.4,None),
      ('HIRO Keywords',SERIF_LATIN,SERIF_CJK,9.15,None,None,HIRO_TEXT,7,7,None,0,0,14.4,None),
      ('HIRO Original',SERIF_LATIN,'Noto Serif CJK JP',9.6,None,None,HIRO_GRAY,8,8,None,4,2,15.2,None),
      ('HIRO Translation',SERIF_LATIN,SERIF_CJK,8.9,None,None,HIRO_LIGHT,8,8,None,0,5,14.3,None),
      ('HIRO Scene Label',SANS_LATIN,SANS_CJK,6.6,None,None,HIRO_LIGHT,None,None,None,14,2,8,None),
      ('HIRO Scene Title',SERIF_LATIN,SERIF_CJK,12.2,True,None,HIRO_INK,None,None,None,0,5,18,None),
      ('HIRO Scene Body',SERIF_LATIN,SERIF_CJK,9.8,None,None,HIRO_TEXT,5,5,7.3,0,8,15.7,None),
      ('HIRO Dialogue',SERIF_LATIN,SERIF_CJK,9.7,None,None,HIRO_TEXT,8,8,None,0,5,15.5,None),
      ('HIRO Code', 'DejaVu Sans Mono','Noto Sans Mono CJK SC',8.6,None,None,HIRO_TEXT,5,5,None,4,4,12,None),
      ('HIRO Reference',SERIF_LATIN,SERIF_CJK,8.45,None,None,HIRO_TEXT,None,None,None,0,3.6,12.55,None),
    ]
    for s in specs: add_pstyle(doc,*s)


def setup_doc():
    doc = Document(); setup_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.left_margin = Mm(27.5); sec.right_margin = Mm(27.5)
    sec.top_margin = Mm(23); sec.bottom_margin = Mm(25)
    sec.header_distance = Mm(7); sec.footer_distance = Mm(12)
    sec.different_first_page_header_footer = True
    # first-page header/footer intentionally blank: TeX feature titlepage and symposium/essay use empty style.
    h = sec.header.paragraphs[0]
    h.paragraph_format.space_after = Pt(0)
    r = h.add_run('HIRO 2026  /  ' + META['paper_type'].upper())
    set_run_font(r, SANS_LATIN, SANS_CJK, 6.9, color=HIRO_LIGHT)
    r = h.add_run('\t' + META['short_title'])
    set_run_font(r, SANS_LATIN, SANS_CJK, 7.0, color=HIRO_LIGHT)
    pPr = h._p.get_or_add_pPr(); tabs=OxmlElement('w:tabs'); t=OxmlElement('w:tab')
    t.set(qn('w:val'),'right'); t.set(qn('w:pos'),'8780'); tabs.append(t); pPr.append(tabs)
    add_page_field(sec.footer.paragraphs[0])
    return doc


def add_spacer(doc, mm):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Mm(mm); p.paragraph_format.line_spacing = Pt(1)
    return p


def add_feature_front(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Mm(8.7); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(str(LOGO), width=Mm(142.6))
    p = doc.add_paragraph(META['kicker'].upper(), style='AT Feature Kicker')
    p.paragraph_format.space_before = Mm(3.5); p.paragraph_format.space_after = Mm(3.8)
    p = doc.add_paragraph(META['title'], style='AT Feature Title')
    p.paragraph_format.space_after = Mm(2.5)
    p = doc.add_paragraph(META['subtitle'], style='AT Feature Subtitle')
    p.paragraph_format.space_after = Mm(2.5)
    doc.add_paragraph(META['english_title'], style='AT Feature English')
    add_spacer(doc, 112)
    doc.add_paragraph(META['author'], style='AT Feature Author')
    p=doc.add_paragraph(META['affiliation'], style='AT Feature Affiliation'); p.paragraph_format.space_before=Mm(1.2)
    p=doc.add_paragraph(META['contact'], style='AT Feature Contact'); p.paragraph_format.space_before=Mm(0.8)
    p=doc.add_paragraph('2026', style='AT Feature Author'); p.paragraph_format.space_before=Mm(2.6)
    p=doc.add_paragraph(META['paper_type']+' / '+META['paper_id'], style='AT Feature Meta'); p.paragraph_format.space_before=Mm(1.2)
    p=doc.add_paragraph(META['author_note'], style='AT Feature Note'); p.paragraph_format.space_before=Mm(2.2)
    doc.add_page_break()


def add_symposium_masthead(doc):
    table = doc.add_table(rows=1, cols=2); table.autofit=False; remove_table_borders(table)
    c0,c1 = table.rows[0].cells
    set_cell_width(c0,54.5); set_cell_width(c1,100.5)
    for c in (c0,c1): set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c0.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE; p.paragraph_format.line_spacing=1.0; p.add_run().add_picture(str(LOGO), width=Mm(39))
    p=c1.paragraphs[0]; p.style='AT Pub Meta'; p.add_run('HIRO2026 / 篠泽广研讨会')
    c1.add_paragraph('SHINOSAWA HIRO / ACGN STUDIES / LONGFORM', style='AT Pub Meta')
    c1.add_paragraph('2026  ·  第三方非官方支援项目', style='AT Pub Meta')
    p=doc.add_paragraph(); p.paragraph_format.space_before=Mm(4.2); p.paragraph_format.space_after=Pt(0); paragraph_bottom_border(p, AT_RULE, '3','0')


def add_author_table(doc, layout):
    t = doc.add_table(rows=1, cols=2); t.autofit=False; remove_table_borders(t)
    c0,c1=t.rows[0].cells
    set_cell_width(c0,106); set_cell_width(c1,49)
    for c in (c0,c1): set_cell_margins(c)
    c0.paragraphs[0].style = 'AT Essay Author' if layout=='essay' else 'AT Author'
    c0.paragraphs[0].add_run(META['author'])
    c0.add_paragraph(META['affiliation'], style='AT Essay Affiliation' if layout=='essay' else 'AT Affiliation')
    if layout=='symposium': c0.add_paragraph(META['contact'], style='AT Contact')
    c1.paragraphs[0].style='AT Document Meta'; c1.paragraphs[0].add_run((META['paper_type']+' / '+META['paper_id']).upper())


def add_compact_front(doc, layout):
    add_symposium_masthead(doc)
    add_spacer(doc, 10 if layout=='essay' else 12.5)
    p=doc.add_paragraph(META['kicker'].upper(), style='AT Kicker'); p.paragraph_format.space_after=Mm(3.8 if layout=='symposium' else 3.5)
    p=doc.add_paragraph(META['title'], style='AT Title'); p.paragraph_format.space_after=Mm(4.2 if layout=='symposium' else 3.8)
    if META['subtitle']:
        p=doc.add_paragraph(META['subtitle'], style='AT Subtitle'); p.paragraph_format.space_after=Mm(4.0 if layout=='symposium' else 0)
    if layout=='symposium':
        p=doc.add_paragraph(META['english_title'], style='AT English'); p.paragraph_format.space_after=Mm(8.5)
    else:
        p=doc.add_paragraph(META['deck'], style='AT Deck'); p.paragraph_format.space_before=Mm(5.5); p.paragraph_format.space_after=Mm(6.5)
    add_author_table(doc, layout)
    p=doc.add_paragraph(META['author_note'], style='AT Note'); p.paragraph_format.space_before=Mm(1.8 if layout=='essay' else 2.0)
    p.paragraph_format.space_after=Mm(7 if layout=='essay' else 8.5)


def add_abstract_keywords(doc):
    p=doc.add_paragraph(style='HIRO Abstract')
    r=p.add_run('摘　要'); set_run_font(r,SANS_LATIN,SANS_CJK,8.7,bold=True,color=HIRO_INK)
    p.add_run('　'+ABSTRACT)
    p.paragraph_format.space_after=Mm(2)
    p=doc.add_paragraph(style='HIRO Keywords')
    r=p.add_run('关键词'); set_run_font(r,SANS_LATIN,SANS_CJK,8.6,bold=True,color=HIRO_INK)
    p.add_run('　'+KEYWORDS)
    p.paragraph_format.space_after=Mm(5.5)
    paragraph_bottom_border(p, HIRO_RULE, '3','0')
    # In TeX the rule is followed by 7.5 mm.
    add_spacer(doc, 2.0)


def add_scene(doc):
    p=doc.add_paragraph(style='HIRO Scene Label')
    paragraph_bottom_border(p, HIRO_RULE, '2','0')
    r=p.add_run('SCENE'); set_run_font(r,SANS_LATIN,SANS_CJK,6.6,color=HIRO_LIGHT)
    r=p.add_run('\tDEMO / SCENE 01'); set_run_font(r,SANS_LATIN,SANS_CJK,6.6,color=HIRO_LIGHT)
    pPr=p._p.get_or_add_pPr(); tabs=OxmlElement('w:tabs'); t=OxmlElement('w:tab'); t.set(qn('w:val'),'right'); t.set(qn('w:pos'),'7600'); tabs.append(t); pPr.append(tabs)
    doc.add_paragraph('示例场景', style='HIRO Scene Title')
    doc.add_paragraph('场景环境用于组织剧情材料或场景化转述，section 继续承担文章的论证层级。', style='HIRO Scene Body')


def add_dialogue(doc):
    p=doc.add_paragraph(style='HIRO Dialogue')
    r=p.add_run('角色 A'); set_run_font(r,SANS_LATIN,SANS_CJK,8,bold=True,color=HIRO_LIGHT)
    p.add_run('　这里展示对话环境的第一段文本。')
    p=doc.add_paragraph(style='HIRO Dialogue')
    r=p.add_run('角色 B'); set_run_font(r,SANS_LATIN,SANS_CJK,8,bold=True,color=HIRO_LIGHT)
    p.add_run('　作者可以把角色对白与分析正文放在不同的视觉层级中。')


def add_body(doc):
    doc.add_paragraph('一、多语种与叙事环境', style='Heading 1')
    doc.add_paragraph('中文姓名“篠泽广”和日文姓名“篠澤 広”可以在同一篇长文中稳定排版。')
    doc.add_paragraph('篠澤 広――日本語組版の確認用テキスト。', style='HIRO Original')
    doc.add_paragraph('篠泽广——用于检查中文译文排版的示例文本。', style='HIRO Translation')
    add_scene(doc); add_dialogue(doc)

    doc.add_paragraph('二、跨媒介参考文献', style='Heading 1')
    doc.add_paragraph('本节使用公开官方元数据展示 GAME、CHARACTER FILE、COMM 与 MUSIC。正式著录由 GB/T 7714 体系承担，HIRO2026 profile 负责专题视觉与 media tag 呈现。')
    doc.add_paragraph('游戏本体可以直接引用为《学園アイドルマスター》[1]，角色资料引用篠澤広官方档案[2]。作品内部剧情节点使用 postnote 保存精确 Locator，例如：')
    doc.add_paragraph(r'\\cite[STEP1 / Episode 8]{hiro_commu_step1_08}', style='HIRO Code')
    doc.add_paragraph('实际渲染为剧情节点[3]。音乐材料可以并引官方曲目页与实体单曲页[4–5]。')
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); paragraph_bottom_border(p,HIRO_RULE,'3','0')
    p=doc.add_paragraph(); r=p.add_run('ATX-ACGN-REF 0.4 / GB/T 7714'); set_run_font(r,SANS_LATIN,SANS_CJK,6.8,bold=True,color=HIRO_LIGHT)
    p.paragraph_format.space_after=Pt(2)
    doc.add_paragraph('GAME  游戏本体　 COMM · STEP4/EP37  游戏剧情　 ANIME · EP01  动画单集　 MANGA · CH01  漫画单话', style='HIRO Reference')
    doc.add_paragraph('MUSIC  音乐　 LIVE  演出　 FAN LOCATOR  定位资料', style='HIRO Reference')
    p=doc.add_paragraph(); paragraph_bottom_border(p,HIRO_RULE,'3','0')
    doc.add_paragraph('参考文献', style='Heading 1')
    for txt in [
        '[1] Bandai Namco Entertainment Inc. 学園アイドルマスター [GAME]. 2024.',
        '[2] 学園アイドルマスター. 篠澤 広 CHARACTER FILE [CHARACTER FILE].',
        '[3] 学園アイドルマスター. STEP1 / Episode 8 [COMM].',
        '[4] 学園アイドルマスター. 光景 [MUSIC].',
        '[5] 篠澤 広 1st Single [MUSIC].',
    ]: doc.add_paragraph(txt, style='HIRO Reference')


def build(layout):
    doc=setup_doc()
    if layout=='feature':
        add_feature_front(doc)
        add_abstract_keywords(doc)
        doc.add_page_break()
    else:
        add_compact_front(doc, layout)
        add_abstract_keywords(doc)
        doc.add_page_break()
    add_body(doc)
    path=OUT/f'hiro2026-{layout}.docx'; doc.save(path); print(path)

if __name__=='__main__':
    for layout in ('feature','symposium','essay'): build(layout)
